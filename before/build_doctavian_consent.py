"""Build the synthetic Doctavian consent template as a deterministic DOCX.

The source document deliberately contains Doctavian merge fields and elements.
It is a treatment-party consent record, never a legal or clinical safety verdict.
"""

from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "before" / "doctavian" / "tx-neurotoxin-consent-v1.docx"
BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK_BLUE = RGBColor(0x17, 0x35, 0x4D)
SLATE = RGBColor(0x33, 0x41, 0x55)
MUTED = RGBColor(0x5F, 0x6B, 0x78)
TEAL = RGBColor(0x0D, 0x76, 0x7A)
PALE = "EAF4F4"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def canonicalize_docx(path: Path) -> None:
    """Normalize ZIP metadata so identical template inputs produce identical bytes."""
    temporary = path.with_suffix(".canonical.tmp")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, date_time=(2026, 8, 18, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = source.getinfo(name).external_attr
            target.writestr(info, source.read(name))
    temporary.replace(path)

def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    properties.append(element)


def set_cell_margins(cell, *, top: int = 120, start: int = 140, bottom: int = 120, end: int = 140) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    """Keep one Doctavian token in one Word run."""
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_label_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = SLATE
    add_field(paragraph, value, color=DARK_BLUE)


def add_signature_anchor(document: Document, role: str, anchor: str, signer_index: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"{role} signature")
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    line = document.add_paragraph("Signature is placed here by Doctavian after identity and intent are confirmed.")
    line.paragraph_format.space_after = Pt(1)
    line.runs[0].font.size = Pt(9)
    line.runs[0].font.color.rgb = MUTED
    anchor_paragraph = document.add_paragraph()
    anchor_paragraph.paragraph_format.space_after = Pt(1)
    anchor_run = anchor_paragraph.add_run(anchor)
    anchor_run.font.color.rgb = WHITE
    anchor_run.font.size = Pt(6)
    signed = document.add_paragraph()
    add_field(signed, f"Signed by: {{!Signatures[{signer_index}].Name}}    Date: {{!Signatures[{signer_index}].SignedAt}}")


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = SLATE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 16, BLUE, 16, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    evidence = document.styles.add_style("Evidence", 1)
    evidence.font.name = "Consolas"
    evidence.font.size = Pt(8.5)
    evidence.font.color.rgb = DARK_BLUE
    evidence.paragraph_format.space_after = Pt(3)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title = paragraph.add_run("BEFORE  /  TEXAS NEUROTOXIN CONSENT")
    title.bold = True
    title.font.color.rgb = TEAL
    title.font.size = Pt(9)
    meta = paragraph.add_run("\nSYNTHETIC DEMONSTRATION  •  TEMPLATE V1")
    meta.font.color.rgb = MUTED
    meta.font.size = Pt(8)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("BEFORE • Treatment-party consent • Not a legality or safety certification")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def build() -> Path:
    document = Document()
    configure_page(document)
    configure_styles(document)
    document.core_properties.title = "BEFORE Texas Neurotoxin Consent — Synthetic Template"
    document.core_properties.subject = "Doctavian conditional consent and treatment-party signatures"
    document.core_properties.author = "BEFORE synthetic demo"
    document.core_properties.keywords = "synthetic, Texas, neurotoxin, consent, Doctavian"
    document.core_properties.comments = "Contains synthetic demonstration fields only."

    title = document.add_paragraph(style="Title")
    title.add_run("Consent to a neurotoxin procedure")
    subtitle = document.add_paragraph("Texas • synthetic demonstration encounter")
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.runs[0].font.color.rgb = MUTED

    notice = document.add_table(rows=1, cols=1)
    notice.autofit = False
    notice.columns[0].width = Inches(6.5)
    cell = notice.cell(0, 0)
    set_cell_margins(cell)
    shade(cell, PALE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("SYNTHETIC DEMONSTRATION — NOT CLINICAL CONSENT")
    run.bold = True
    run.font.color.rgb = TEAL
    paragraph.add_run("\nThis document records disclosure and treatment-party signatures. It does not determine legality, certify safety, prescribe a dose, or replace professional judgment.")

    document.add_heading("Encounter", level=1)
    add_label_value(document, "Encounter ID", "{!Encounter[0].EncounterId}")
    add_label_value(document, "Patient", "{!Encounter[0].PatientDisplayName}")
    add_label_value(document, "Procedure", "{!Encounter[0].ProcedureDisplayName}")
    add_label_value(document, "Performer", "{!Encounter[0].PerformerDisplayName}")
    add_label_value(document, "Scheduled", "{!Encounter[0].ScheduledOn}")

    document.add_heading("What I am considering", level=1)
    document.add_paragraph("I am considering an elective neurotoxin injection. The expected effect is temporary. Results vary, additional treatment may be discussed, and choosing no treatment is an alternative.")
    document.add_paragraph("I had an opportunity to ask questions. My care team must explain material risks, alternatives, aftercare, and who will perform the procedure in language I can understand.")

    conditional = document.add_paragraph()
    add_field(conditional, '<mdoc:paragraph name="delegatedAuthority" hidden="{!$Encounter[0].AuthorityPathway != \'DELEGATED\'}">')
    conditional.add_run("The proposed performer is acting through a documented delegation pathway. I was told the performer's role and how the delegating clinician is available for supervision and escalation.")
    add_field(conditional, "</mdoc:paragraph>")

    patient_hold = document.add_paragraph()
    add_field(patient_hold, '<mdoc:paragraph name="patientFlagReview" hidden="{!$Encounter[0].PatientFlagReviewRequired != true}">')
    patient_hold.add_run("A patient-specific flag requires named human review before the procedure may progress. This consent does not clear or resolve that hold.")
    add_field(patient_hold, "</mdoc:paragraph>")

    document.add_page_break()
    document.add_heading("Required disclosures", level=1)
    count = document.add_paragraph(style="Evidence")
    add_field(count, "Disclosure count: {!$count(Encounter[0].RequiredDisclosures)}")

    start = document.add_paragraph()
    add_field(start, '<mdoc:repeater name="requiredDisclosures" value="Encounter[0].RequiredDisclosures" variable="disclosure">')
    disclosure_heading = document.add_paragraph(style="Heading 2")
    add_field(disclosure_heading, "{!#disclosure#.Title}")
    body = document.add_paragraph()
    add_field(body, "{!#disclosure#.PlainLanguage}")
    citation = document.add_paragraph(style="Evidence")
    add_field(citation, "SOURCE  {!#disclosure#.Citation}")
    end = document.add_paragraph()
    add_field(end, "</mdoc:repeater>")

    document.add_heading("Frozen rule context", level=1)
    document.add_paragraph("The following identifiers preserve which deterministic rules and cited sources were shown when this record was created. They support later explanation; they are not legal conclusions.")
    snapshot = document.add_paragraph(style="Evidence")
    add_field(snapshot, "RULE  {!Encounter[0].RuleId}")
    snapshot = document.add_paragraph(style="Evidence")
    add_field(snapshot, "SNAPSHOT SHA-256  {!Encounter[0].RuleSnapshotSha256}")

    document.add_heading("Acknowledgements", level=1)
    for item in (
        "I received an explanation of expected effects, material risks, reasonable alternatives, and aftercare.",
        "I understand that signing records my consent; it does not guarantee an outcome or certify safety.",
        "I understand that a safety hold or human-review task may still prevent the procedure from progressing.",
        "I understand that the medical director attestation, if required, is a separate Foxit signing event.",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_page_break()
    document.add_heading("Treatment-party signatures", level=1)
    document.add_paragraph("Both signatures are required. Doctavian records the patient and injector acknowledgements. A medical director must not sign in either of these roles.")
    add_signature_anchor(document, "Patient", "_SIG_PATIENT_", 0)
    add_signature_anchor(document, "Injector", "_SIG_INJECTOR_", 1)

    document.add_heading("Machine-checkable context", level=2)
    for label, token in (
        ("TEMPLATE", "TX-NEUROTOXIN-CONSENT-1"),
        ("ENCOUNTER", "{!Encounter[0].EncounterId}"),
        ("RULE", "{!Encounter[0].RuleId}"),
        ("SNAPSHOT", "{!Encounter[0].RuleSnapshotSha256}"),
    ):
        paragraph = document.add_paragraph(style="Evidence")
        add_field(paragraph, f"{label}  {token}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    canonicalize_docx(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
