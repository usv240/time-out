"""Generate the consent document through Doctavian's API, for real.

What works and what does not
----------------------------
Established by generating documents and reading the rendered PDF back, rather than
trusting HTTP 200 — several unsupported expressions return 200 and render blank, which
is a worse failure than an error:

    {!Path}                          renders
    {!$count(array)}                 renders — calculation on the platform
    {!array[0].Field}                renders — indexing into a repeated record
    ternary / $if(...)               returns 200 and renders EMPTY
    <mdoc:repeater> / <mdoc:paragraph>   PROCESS_MARKUP_ELEMENT_FAILED

The markup elements are authored through Doctavian's Office add-in and cannot be
produced programmatically, so this template uses expressions, a platform-side count,
and indexed access to the cited-disclosure records.

The branch stays where every other decision in this product lives: the Gate resolves
the authority pathway from cited rules and passes the resolved sentence in as data. A
document template is the wrong place to decide who may perform a procedure.

Two contract details that cost real time, both verified by controlled runs:
  * the uploaded JSON must be wrapped in one root `data` object. Omit it and the call
    fails with `TEMPLATE_READ_FAILED` — "check the template format" — which names the
    wrong artifact entirely. Scalar types do not matter; the envelope does.
  * `request/create` takes `dataSourceGuid`, not the documented `dataGuid`

    python -m before.doctavian_generate
"""
from __future__ import annotations

import io
import json
import os
import time
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "before" / "doctavian" / "consent-data.synthetic.json"
OUT = ROOT / "before" / "doctavian" / "generated"
CACHE = ROOT / ".cache" / "doctavian"
BASE = "https://demo.api.doctavian.com"
MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# The Gate decides this; the template only renders it.
AUTHORITY_SENTENCE = {
    "DELEGATED": ("The proposed performer is not the delegating physician. This procedure "
                  "proceeds under documented delegation, a signed protocol, a patient-specific "
                  "order, and an available supervisor."),
    "DIRECT": ("The proposed performer is the physician. No delegation is asserted for this "
               "procedure."),
}


def _headers(content_type: str | None = None, storage: str | None = None) -> dict[str, str]:
    from before.doctavian_auth import bearer
    h = {"x-api-key": os.environ["DOCTAVIAN_API_KEY"],
         "Authorization": f"Bearer {bearer()}",
         "X-Origin": os.environ.get("DOCTAVIAN_ORIGIN", "https://app.mavenmule.com")}
    if content_type:
        h["Content-Type"] = content_type
    if storage:
        h["X-Storage-Type"] = storage
    return h


def _stringify(value: Any) -> Any:
    """Every scalar leaf becomes a string; the object and array structure is kept.

    Not required — a controlled run with the same template and raw booleans and numbers
    also succeeds. What *is* required is the root `data` wrapper below. This is kept
    because it makes rendering predictable: the platform formats a string exactly as
    given, so a date or a score cannot be re-formatted into something a clinician did
    not write.
    """
    if isinstance(value, dict):
        return {k: _stringify(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_stringify(v) for v in value]
    if isinstance(value, bool):
        return "true" if value else "false"
    if value is None:
        return ""
    return str(value)


def build_template(disclosure_count: int) -> bytes:
    """The consent template, using only expressions Doctavian actually renders."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_paragraph("Consent to a neurotoxin procedure").runs[0].font.size = Pt(18)
    doc.add_paragraph("Texas · synthetic demonstration encounter")

    doc.add_paragraph("Encounter").runs[0].bold = True
    for label, expr in (("Encounter", "EncounterId"), ("Patient", "PatientDisplayName"),
                        ("Procedure", "ProcedureDisplayName"), ("Performer", "PerformerDisplayName"),
                        ("Scheduled", "ScheduledOn")):
        doc.add_paragraph(f"{label}: {{!Encounter[0].{expr}}}")

    doc.add_paragraph("Who is performing this procedure").runs[0].bold = True
    doc.add_paragraph("Authority pathway: {!Encounter[0].AuthorityPathway}")
    # Resolved by the Gate from cited rules, not decided in this document.
    doc.add_paragraph("{!Encounter[0].AuthorityStatement}")

    doc.add_paragraph("What I am considering").runs[0].bold = True
    doc.add_paragraph("I am considering an elective neurotoxin injection. The expected effect is "
                      "temporary. Results vary, further treatment may be discussed, and choosing "
                      "no treatment is an alternative.")

    doc.add_paragraph("Required disclosures").runs[0].bold = True
    # $count is evaluated by Doctavian, not precomputed by us.
    doc.add_paragraph("Disclosures cited for this encounter: "
                      "{!$count(Encounter[0].RequiredDisclosures)}")
    for i in range(disclosure_count):
        doc.add_paragraph(f"{i + 1}. {{!Encounter[0].RequiredDisclosures[{i}].Title}}")
        doc.add_paragraph(f"   {{!Encounter[0].RequiredDisclosures[{i}].PlainLanguage}}")
        doc.add_paragraph(f"   Source: {{!Encounter[0].RequiredDisclosures[{i}].Citation}}")

    doc.add_paragraph("Frozen rule context").runs[0].bold = True
    doc.add_paragraph("Ruleset: {!Encounter[0].RuleId}")
    doc.add_paragraph("Snapshot: {!Encounter[0].RuleSnapshotSha256}")

    doc.add_paragraph("Signatures").runs[0].bold = True
    doc.add_paragraph("Patient: {!Encounter[0].PatientDisplayName} ______________________")
    doc.add_paragraph("Injector: {!Encounter[0].PerformerDisplayName} ______________________")
    doc.add_paragraph("This record is synthetic. It contains no real patient, clinic, licence, "
                      "lot, or signature.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload(path: str, filename: str, blob: bytes, mime: str, storage: str) -> str:
    import requests
    r = requests.post(f"{BASE}{path}", headers=_headers(storage=storage),
                      files={"file": (filename, blob, mime)}, timeout=120)
    r.raise_for_status()
    return r.json()["result"]["data"]["files"][0]["id"]


def generate(attempts: int = 5) -> dict[str, Any]:
    import requests

    encounter = json.loads(DATA.read_text(encoding="utf-8"))["Encounter"][0]
    encounter = dict(encounter)
    encounter["AuthorityStatement"] = AUTHORITY_SENTENCE.get(
        encounter.get("AuthorityPathway", ""), AUTHORITY_SENTENCE["DIRECT"])
    payload = {"data": _stringify({"Encounter": [encounter]})}
    template = build_template(len(encounter.get("RequiredDisclosures", [])))

    last: dict[str, Any] = {}
    for attempt in range(attempts):
        # Uploaded files are only briefly visible to the generator, and a fresh upload
        # occasionally reports as missing on the very next call. Re-upload and retry.
        tpl = _upload("/v1/documents/template/upload", "tx-neurotoxin-consent.docx",
                      template, MIME, "document-template")
        dat = _upload("/v1/documents/data/upload", "consent-data.json",
                      json.dumps(payload, separators=(",", ":")).encode(),
                      "application/json", "document-data")
        body = {
            "externalContext": {"id": encounter["EncounterId"]},
            "template": {"name": "tx-neurotoxin-consent.docx", "urn": tpl,
                         "fileFormat": "docx", "loadMethod": "Storage", "options": {}},
            "data": {"loadMethod": "Storage", "urn": dat},
            "document": {"timezone": "America/Chicago", "locale": "en_US_POSIX",
                         "name": f"consent-{encounter['EncounterId']}", "fileFormat": "pdf",
                         "deliveryMethod": "Storage", "path": "root", "options": {}},
        }
        r = requests.post(f"{BASE}/v1/documents/document/generate",
                          headers=_headers("application/json"), json=body, timeout=240)
        last = r.json()
        if r.status_code < 400:
            urn = last["result"]["data"]["document"]["urn"]
            pdf = requests.get(f"{BASE}/v1/documents/document/{urn}/download",
                               headers=_headers(), timeout=180)
            pdf.raise_for_status()
            if not pdf.content.startswith(b"%PDF-"):
                raise RuntimeError("Doctavian returned something that is not a PDF.")
            OUT.mkdir(parents=True, exist_ok=True)
            out = OUT / "tx-neurotoxin-consent.pdf"
            out.write_bytes(pdf.content)
            record = {
                "document_urn": urn,
                "bytes": len(pdf.content),
                "encounter_id": encounter["EncounterId"],
                "authority_pathway": encounter["AuthorityPathway"],
                "disclosures_cited": len(encounter.get("RequiredDisclosures", [])),
                "consumption": last.get("consumption"),
                "path": str(out.relative_to(ROOT)).replace("\\", "/"),
                "attempts": attempt + 1,
                "boundary": ("Doctavian rendered this document. The authority pathway was "
                             "resolved by the Gate from cited rules and passed in as data; a "
                             "document template is the wrong place to decide who may perform "
                             "a procedure."),
            }
            CACHE.mkdir(parents=True, exist_ok=True)
            (CACHE / "generate-live.json").write_text(json.dumps(record, indent=2), encoding="utf-8")
            return record
        codes = [e.get("code") for e in last.get("error", {}).get("innerErrors", [])]
        if "FILE_MISSING_FROM_STORAGE" not in str(codes):
            raise RuntimeError(f"Doctavian generate failed: {codes}")
        time.sleep(2)
    raise RuntimeError(f"Doctavian generate never became visible in storage: {last}")


if __name__ == "__main__":
    print(json.dumps(generate(), indent=2))
